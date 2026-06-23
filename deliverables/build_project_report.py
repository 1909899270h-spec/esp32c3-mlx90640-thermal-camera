"""Build the full project report for the ESP32-C3 + MLX90640 course project.

This report is longer and more complete than the earlier checklist-oriented
submission report. It combines the project development process, hardware
wiring, software architecture, debugging logic, current results, and follow-up
improvements into one formal course report maintained by rdfrdream.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from build_development_process_doc import prepare_assets


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "基于ESP32-C3与MLX90640的热成像相机系统项目报告.docx"

FONT_NAME = "微软雅黑"
MONO_FONT = "Consolas"
ACCENT = RGBColor(31, 78, 121)
ACCENT_DARK = RGBColor(15, 44, 76)
BODY = RGBColor(38, 45, 54)
MUTED = RGBColor(92, 102, 112)
LIGHT_FILL = "EEF4FA"
NOTE_FILL = "F5F8FB"


def set_run_font(run, size: float = 10.5, bold: bool = False, color: RGBColor = BODY, name: str = FONT_NAME):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color: str = "D7DEE8", size: str = "8"):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def paragraph_border_bottom(paragraph, color: str = "2E74B5", size: str = "14"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.75)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.85)
    section.right_margin = Cm(1.85)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color, before, after in [
        ("Heading 1", 18, ACCENT, 12, 7),
        ("Heading 2", 14, ACCENT, 10, 5),
        ("Heading 3", 12, ACCENT_DARK, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("微机接口与技术期末大作业 | ESP32-C3 + MLX90640 热成像相机系统")
    set_run_font(run, size=9, color=MUTED)

    props = doc.core_properties
    props.author = "rdfrdream"
    props.title = "基于 ESP32-C3 与 MLX90640 的热成像相机系统项目报告"
    props.subject = "微机接口与技术期末大作业项目报告"
    return doc


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = False
    return p


def add_para(doc: Document, text: str, size: float = 10.5, bold: bool = False, color: RGBColor = BODY, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_note(doc: Document, title: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    cell = table.cell(0, 0)
    shade_cell(cell, NOTE_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11.3, bold=True, color=ACCENT_DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.2):
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(8)
    run = cp.add_run(caption)
    set_run_font(run, size=9.2, color=MUTED)


def add_image_grid(doc: Document, items: list[tuple[Path, str]], width: float = 2.95):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "FFFFFF", "0")
    for i in range(0, len(items), 2):
        row = table.add_row()
        for j in range(2):
            cell = row.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if i + j >= len(items):
                continue
            image_path, caption = items[i + j]
            if image_path.exists():
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(image_path), width=Inches(width))
            cp = cell.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cp.add_run(caption)
            set_run_font(r, size=9, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc: Document, headers: tuple[str, ...], rows: list[tuple[str, ...]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        if widths:
            cell.width = Inches(widths[i])
        shade_cell(cell, LIGHT_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.8, bold=True, color=ACCENT_DARK)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            if widths:
                cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_command_block(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "D7DEE8", "6")
    cell = table.cell(0, 0)
    shade_cell(cell, "F6F8FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for line_no, line in enumerate(text.splitlines()):
        if line_no:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, size=9, color=BODY, name=MONO_FONT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("基于 ESP32-C3 与 MLX90640 的\n热成像相机系统项目报告")
    set_run_font(r, size=25, bold=True, color=ACCENT_DARK)
    paragraph_border_bottom(p)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(18)
    r = sub.add_run("微机接口与技术 / 微机原理期末大作业 | rdfrdream")
    set_run_font(r, size=12.5, color=MUTED)

    add_note(
        doc,
        "摘要",
        "本项目围绕低成本热成像检测与网页可视化展开，使用 ESP32-C3 作为主控，连接 MLX90640 32 x 24 红外阵列传感器，通过 I2C 读取温度矩阵，并利用 ESP32-C3 自带 WiFi 建立局域网 Web 服务。浏览器端网页负责热图绘制、色温条显示、数据刷新、演示数据导入、全屏展示和 CSV 导出。项目完成了 PlatformIO 工程、ESP32-C3 固件、SPIFFS 网页资源、本地演示服务器、硬件接线方案、调试接口和文档整理；真实 MLX90640 采集链路已进入实物稳定性验证阶段。本报告重点记录项目背景、总体方案、硬件连接、软件实现、调试过程、现阶段结果和后续改进方向。",
    )

    add_table(
        doc,
        ("项目要素", "说明"),
        [
            ("主控平台", "ESP32-C3 开发板，负责 I2C 采集、WiFi AP、HTTP Server 和数据接口"),
            ("传感器", "MLX90640 红外阵列传感器，32 x 24，共 768 个温度点，I2C 地址 0x33"),
            ("显示方式", "手机或电脑连接 ESP32-C3 热点后，通过浏览器访问网页热图"),
            ("工程环境", "VS Code + PlatformIO + ESP-IDF，支持编译、烧录、SPIFFS 上传和串口监视"),
            ("仓库地址", "https://github.com/1909899270h-spec/esp32c3-mlx90640-thermal-camera"),
        ],
        [1.3, 5.0],
    )
    doc.add_page_break()


def section_background(doc: Document):
    add_heading(doc, "1. 项目背景与设计目标", 1)
    add_para(
        doc,
        "热成像技术能够把物体表面的红外辐射转换为可观察的温度分布图，在电路板发热检测、设备状态观察、实验教学演示和低成本测温装置中具有直观价值。传统热像仪成本较高，且通常作为独立仪器使用，不便于与课程中的单片机接口、无线通信和网页可视化内容结合。ESP32-C3 自带 WiFi/BLE，资源足够运行轻量级 Web Server；MLX90640 提供 32 x 24 阵列温度数据，两者组合后可以构成一个低成本、可展示、可记录的热成像系统。",
    )
    add_para(
        doc,
        "本项目的设计目标不是单纯点亮一个传感器，而是完成从红外阵列采集到网页显示的完整链路：传感器通过 I2C 向 ESP32-C3 提供温度数据；ESP32-C3 进行数据缓存并提供 HTTP 接口；手机或电脑通过 WiFi 连接到 ESP32-C3，使用浏览器查看热图和温度数据；网页端能够记录数据并导出 CSV，便于后续分析。",
    )
    add_bullets(
        doc,
        [
            "教学目标：覆盖 I2C 总线、GPIO、嵌入式编译烧录、串口调试、WiFi AP、HTTP 接口和浏览器可视化。",
            "工程目标：形成可以演示的硬件样机和软件系统，并将源代码、文档、图片证据和运行方式整理到 GitHub 仓库。",
            "展示目标：即使真实传感器临时连接不稳定，也能通过本地演示数据说明网页端和数据链路的完整功能。",
        ],
    )


def section_system_design(doc: Document, assets: dict[str, Path]):
    add_heading(doc, "2. 系统总体方案", 1)
    add_para(
        doc,
        "系统采用“传感器采集 - 主控处理 - 无线访问 - 网页显示 - 数据记录”的分层结构。MLX90640 负责红外阵列测温，ESP32-C3 负责 I2C 通信、温度帧缓存、WiFi 热点和 HTTP 服务，浏览器网页负责热图绘制、按钮交互和 CSV 导出。该结构把底层采集和上层显示分开，便于调试和替换硬件。",
    )
    add_figure(doc, assets["flow"], "图 1 系统总体数据链路", width=6.7)
    add_table(
        doc,
        ("层级", "功能", "本项目实现"),
        [
            ("硬件层", "供电、传感器连接、I2C 信号", "MLX90640 VDD/GND/SDA/SCL 接入 ESP32-C3"),
            ("采集层", "读取红外阵列原始数据并计算温度", "src/main.c 调用 MLX90640 API 完成帧读取"),
            ("网络层", "建立无线网络与网页服务", "ESP32-C3 AP 模式，默认访问 192.168.4.1"),
            ("显示层", "热图、色温条、数据统计", "data/web_script.js 使用 Canvas 绘制热图"),
            ("记录层", "保存时间序列温度数据", "网页端记录帧数据并下载 CSV"),
        ],
        [1.05, 2.0, 3.1],
    )
    add_note(
        doc,
        "为什么选择 ESP32-C3",
        "C51 适合基础控制但资源较少，联网和网页显示困难；STM32 控制能力强、外设丰富，但通常需要额外 WiFi 或以太网模块；ESP32-C3 自带 WiFi/BLE，可以同时承担主控、无线通信和小型 Web Server 三个角色，更适合本项目这种低成本、无线化、易展示的热成像系统。",
    )


def section_hardware(doc: Document, assets: dict[str, Path]):
    add_heading(doc, "3. 硬件设计与接口连接", 1)
    add_para(
        doc,
        "硬件样机使用 ESP32-C3 开发板、MLX90640 裸传感器、面包板、杜邦线、5kΩ 上拉电阻和 10uF 电容搭建。选择面包板的原因是便于在调试阶段快速调整接线，但裸传感器引脚较细，临时插接会带来松动和接触不稳定问题，这也是后续真实采集链路需要重点优化的地方。",
    )
    add_image_grid(
        doc,
        [
            (assets["esp32_front"], "ESP32-C3 开发板正面"),
            (assets["esp32_back"], "ESP32-C3 背面引脚标识"),
            (assets["mlx90640"], "MLX90640 裸传感器"),
            (assets["wiring"], "面包板接线样机"),
        ],
        width=2.75,
    )
    add_heading(doc, "3.1 核心接线关系", 2)
    add_table(
        doc,
        ("MLX90640 引脚", "ESP32-C3 引脚", "作用", "说明"),
        [
            ("VDD", "3.3V", "传感器供电", "避免接 5V，减少烧坏风险"),
            ("GND", "GND", "公共地", "主控和传感器必须共地"),
            ("SDA", "GPIO4", "I2C 数据线", "建议通过 5kΩ 电阻上拉到 3.3V"),
            ("SCL", "GPIO5", "I2C 时钟线", "建议通过 5kΩ 电阻上拉到 3.3V"),
        ],
        [1.25, 1.3, 1.3, 2.35],
    )
    add_figure(doc, assets["i2c_diagram"], "图 2 MLX90640 与 ESP32-C3 的 I2C 连接说明", width=5.5)
    add_note(
        doc,
        "上拉电阻与去耦电容的作用",
        "I2C 的 SDA/SCL 属于开漏或开集电极风格信号，需要上拉电阻把空闲状态维持为高电平；10uF 电容并联在 VDD 与 GND 之间，用于降低供电波动。对于 MLX90640 这种阵列传感器，稳定供电和可靠接触会直接影响 I2C 扫描和帧数据读取。",
    )


def section_environment(doc: Document, assets: dict[str, Path]):
    add_heading(doc, "4. 开发环境与工程结构", 1)
    add_para(
        doc,
        "软件开发使用 VS Code + PlatformIO。PlatformIO 的作用不仅是编译代码，还包括管理 ESP-IDF 框架、开发板参数、串口烧录、串口监视、SPIFFS 文件系统镜像构建和网页资源上传。这样可以把 ESP32-C3 固件和网页前端放在同一个工程中统一管理。",
    )
    add_figure(doc, assets["dev_env"], "图 3 VS Code 与 PlatformIO 工程环境", width=5.8)
    add_heading(doc, "4.1 仓库关键文件", 2)
    add_command_block(
        doc,
        """esp32c3-mlx90640-thermal-camera/
|-- platformio.ini                  # PlatformIO 工程配置
|-- src/main.c                      # ESP32-C3 主程序
|-- data/web_script.js              # 网页前端脚本
|-- lib/MLX90640/                   # MLX90640 API 与算法库
|-- tools/mock_thermal_server.py    # 本地演示服务器
|-- partitions.csv                  # Flash 分区与 SPIFFS 配置
|-- README.md                       # 运行方式与调试说明
`-- deliverables/                   # 报告、PPT、图片和生成脚本""",
    )
    add_heading(doc, "4.2 基本编译与烧录流程", 2)
    add_command_block(
        doc,
        """pio run -e esp32-c3-devkitm-1
pio run -e esp32-c3-devkitm-1 -t upload
pio run -e esp32-c3-devkitm-1 -t uploadfs
pio device monitor -e esp32-c3-devkitm-1""",
    )
    add_para(
        doc,
        "其中 `upload` 用于烧录主程序，`uploadfs` 用于上传 SPIFFS 网页资源。实际调试中，如果烧录停在 `Connecting...`，需要按住 BOOT/FLASH 键，再轻按 RESET/RST，待出现写入信息后松开 BOOT/FLASH。",
    )


def section_software(doc: Document, assets: dict[str, Path]):
    add_heading(doc, "5. 软件架构与关键实现", 1)
    add_para(
        doc,
        "软件部分围绕两条链路展开：第一条是传感器数据链路，即 I2C 初始化、MLX90640 参数读取、温度帧计算和共享缓存；第二条是网页数据链路，即 ESP32-C3 创建 WiFi 热点、启动 HTTP Server、返回 JSON 数据、浏览器周期性刷新并绘制热图。两条链路相互独立又通过温度缓存连接，便于在硬件不稳定时使用本地演示数据验证网页逻辑。",
    )
    add_table(
        doc,
        ("模块", "主要职责", "实现要点"),
        [
            ("I2C 驱动", "配置 GPIO4/GPIO5、400kHz 时钟、读写寄存器", "封装 MLX90640_I2CRead/Write 供 API 调用"),
            ("传感器初始化", "读取 EEPROM、提取校准参数、设置刷新率", "记录初始化状态，方便串口和网页调试"),
            ("采集任务", "循环读取帧数据并计算 768 点温度", "使用缓冲区和互斥量降低并发访问风险"),
            ("HTTP 接口", "向网页返回温度 JSON、状态和 I2C 扫描结果", "便于判断硬件层、总线层和数据层问题"),
            ("前端网页", "绘制 Canvas 热图、色温条、CSV 导出", "支持开始检测、导入数据、全屏和清空记录"),
        ],
        [1.2, 2.0, 3.0],
    )
    add_figure(doc, assets["code_init"], "图 4 MLX90640 初始化与参数读取代码截图", width=5.8)
    add_figure(doc, assets["code_task"], "图 5 温度帧采集任务代码截图", width=5.8)
    add_figure(doc, assets["code_api"], "图 6 HTTP 温度数据接口代码截图", width=5.8)


def section_web_network(doc: Document, assets: dict[str, Path]):
    add_heading(doc, "6. WiFi Web 服务与网页可视化", 1)
    add_para(
        doc,
        "ESP32-C3 在 AP 模式下可以自己建立 WiFi 热点。手机或电脑连接该热点后，实际上进入了 ESP32-C3 建立的局域网，浏览器访问 `http://192.168.4.1` 时，请求会被发送到 ESP32-C3 内部运行的 HTTP Server。若使用 STA 模式连接路由器，则浏览器访问的是路由器分配给 ESP32-C3 的局域网 IP。",
    )
    add_para(
        doc,
        "这些地址属于局域网私有地址，不是公网地址，因此手机切换到蜂窝网络或其他 WiFi 后无法直接访问。这是网络结构原因，不是网页代码错误。若后续需要远程访问，可以通过云端服务器、中转服务、内网穿透或 MQTT/HTTP 上报等方式把数据转发到公网可访问位置。",
    )
    add_figure(doc, assets["wifi_diagram"], "图 7 ESP32-C3 热点、局域网访问与网页服务关系", width=5.8)
    add_heading(doc, "6.1 网页功能设计", 2)
    add_bullets(
        doc,
        [
            "开始检测：从 ESP32-C3 的 `/thermal-data` 接口读取真实传感器数据。",
            "导入数据：使用网页端平滑演示数据，用于硬件不稳定或录制演示时验证可视化链路。",
            "下载 CSV 数据：按照时间顺序导出已记录的温度帧，便于后续分析。",
            "色温条与温度范围：显示当前最低温、最高温和颜色映射关系。",
            "全屏展示：方便课堂答辩或演示视频录制。",
        ],
    )
    add_figure(doc, assets["web_debug"], "图 8 网页热图、色温条、数据记录和交互按钮效果", width=6.4)


def section_debug_result(doc: Document):
    add_heading(doc, "7. 调试过程与实验结果", 1)
    add_para(
        doc,
        "项目调试遵循“先软件、再通信、后传感器”的思路。首先确认 PlatformIO 工程能编译，ESP32-C3 能被电脑识别为串口设备，主程序和 SPIFFS 网页资源能够烧录；随后通过浏览器确认网页能够打开；最后再进入 I2C 扫描、MLX90640 初始化和温度帧读取验证。",
    )
    add_heading(doc, "7.1 三层问题定位方法", 2)
    add_table(
        doc,
        ("层级", "判断方法", "期望现象", "异常时重点检查"),
        [
            ("硬件层", "看供电、串口、3.3V/GND", "ESP32-C3 上电，电脑识别 COM 口", "Type-C 数据线、GND 共地、杜邦线松动"),
            ("总线层", "I2C scan 或串口日志", "扫描到 MLX90640 地址 0x33", "SDA/SCL 是否接反，上拉是否有效"),
            ("数据层", "网页 min/max/热图变化", "温度随热源移动变化", "初始化状态、帧读取状态、JSON 接口"),
        ],
        [0.9, 1.55, 1.75, 2.05],
    )
    add_heading(doc, "7.2 当前验证结果", 2)
    add_bullets(
        doc,
        [
            "PlatformIO 工程可以完成主固件编译，ESP32-C3 固件和 SPIFFS 网页文件系统均已验证过构建流程。",
            "网页端可以实现热图显示、色温条、开始检测、导入数据、全屏展示、清空记录和 CSV 下载等功能。",
            "本地 Python 演示服务器能够生成平滑移动的热源数据，用于验证网页端热图和数据交互逻辑。",
            "真实 MLX90640 采集链路已完成接线方案和软件接口，但受裸传感器、面包板、杜邦线松动影响，连续采集稳定性仍在验证中。",
        ],
    )
    add_note(
        doc,
        "关于实地验证的说明",
        "在课堂实地验证要求下，本项目已经完成实物搭建、固件烧录、网页访问和演示数据展示。由于 MLX90640 裸传感器与面包板连接不够稳定，真实温度帧在现场无法长时间稳定输出，因此演示时采用前期实验视频和网页端演示数据补充说明系统流程。该问题主要集中在硬件连接稳定性，不是 Web 界面、CSV 导出或 ESP32-C3 网页服务逻辑错误。",
    )


def section_discussion_future(doc: Document):
    add_heading(doc, "8. 问题分析与改进方向", 1)
    add_para(
        doc,
        "从系统完整性看，项目的软件链路已经较完整：ESP32-C3 可以作为服务器提供网页和数据接口，网页端也能够完成热图、色温条和记录功能。现阶段最大的限制是硬件层可靠性，尤其是裸 MLX90640 引脚接触、面包板空间不足、杜邦线松动和 I2C 信号质量。要把系统从演示样机推进到稳定产品，需要优先解决硬件固定和信号完整性。",
    )
    add_heading(doc, "8.1 硬件优化", 2)
    add_bullets(
        doc,
        [
            "使用 MLX90640 转接板或自行设计 PCB，避免裸引脚直接插线导致的接触不良。",
            "缩短 SDA/SCL 线长，保留 5kΩ 上拉电阻和 10uF 去耦电容，必要时重新评估 I2C 速率。",
            "加入外壳和传感器固定结构，保证镜头方向、接线位置和演示时的机械稳定性。",
        ],
    )
    add_heading(doc, "8.2 算法与校准优化", 2)
    add_bullets(
        doc,
        [
            "加入异常点滤波、帧间平滑和热点追踪，使 32 x 24 低分辨率数据在网页端显示更自然。",
            "建立温度校准流程，例如使用已知温度参考源做单点或两点校准，对偏移和增益进行补偿。",
            "根据被测物体材料调整发射率参数，提高不同表面材料下的温度估计一致性。",
        ],
    )
    add_heading(doc, "8.3 功能拓展", 2)
    add_bullets(
        doc,
        [
            "加入温度极值触发 LED 或蜂鸣器报警，用于演示阈值检测和自动响应。",
            "通过云端服务器、内网穿透或数据上报接口，使网页访问不局限于 ESP32-C3 局域网。",
            "结合红外识别算法，扩展为电路板发热检测、设备异常识别或红外目标自动检测装置。",
        ],
    )


def section_summary(doc: Document):
    add_heading(doc, "9. 总结", 1)
    add_para(
        doc,
        "本项目完成了基于 ESP32-C3 与 MLX90640 的热成像相机系统设计与工程整理。系统将 I2C 红外阵列采集、ESP32-C3 WiFi Web 服务、浏览器热图可视化和 CSV 数据记录整合到同一工程中，体现了微机接口技术从底层硬件连接到上层应用展示的完整流程。虽然真实传感器连续采集还受到临时接线稳定性的限制，但项目已经形成可编译、可烧录、可网页演示、可记录数据、可开源提交的完整软件与文档体系。",
    )
    add_para(
        doc,
        "后续若完成传感器转接板或 PCB 固定，并加入校准、滤波、报警和远程访问功能，该系统可以进一步发展为面向实验教学、电路板发热检测或红外异常识别的小型热成像平台。",
    )
    add_heading(doc, "附录：提交材料", 1)
    add_table(
        doc,
        ("材料", "位置或说明"),
        [
            ("完整源码", "src、data、lib、tools、platformio.ini、partitions.csv"),
            ("运行说明", "README.md，包含接线、编译、烧录、网页访问和调试接口"),
            ("项目报告", "deliverables/基于ESP32-C3与MLX90640的热成像相机系统项目报告.pdf"),
            ("开发流程文档", "deliverables/基于ESP32-C3与MLX90640的热成像相机系统开发流程.pdf"),
            ("GitHub 仓库", "https://github.com/1909899270h-spec/esp32c3-mlx90640-thermal-camera"),
        ],
        [1.35, 4.9],
    )


def build_report():
    assets = prepare_assets()
    doc = setup_doc()
    add_title_page(doc)
    section_background(doc)
    section_system_design(doc, assets)
    section_hardware(doc, assets)
    section_environment(doc, assets)
    section_software(doc, assets)
    section_web_network(doc, assets)
    section_debug_result(doc)
    section_discussion_future(doc)
    section_summary(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_report()
