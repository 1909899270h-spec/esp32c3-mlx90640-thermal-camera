"""Build the concise submission report for the course checklist.

The report is a separate deliverable from the development-process document.
It is intentionally kept under 5000 Chinese characters and written in the
name of rdfrdream.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from build_development_process_doc import prepare_assets


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "基于ESP32-C3与MLX90640的热成像相机系统项目报告_提交版.docx"

FONT_NAME = "微软雅黑"
ACCENT = RGBColor(31, 78, 121)
ACCENT_DARK = RGBColor(15, 44, 76)
MUTED = RGBColor(92, 102, 112)
BODY = RGBColor(38, 45, 54)
LIGHT_FILL = "EEF4FA"
NOTE_FILL = "F5F8FB"


def set_run_font(run, size=10.5, bold=False, color=BODY, name=FONT_NAME):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color: str = "D7DEE8", size: str = "8"):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def paragraph_border_bottom(paragraph, color: str = "2E74B5", size: str = "12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Heading 1", 18, ACCENT),
        ("Heading 2", 14, ACCENT),
        ("Heading 3", 12, ACCENT_DARK),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("微机接口与技术期末大作业 | ESP32-C3 + MLX90640 热成像相机系统")
    set_run_font(run, size=9, color=MUTED)

    props = doc.core_properties
    props.author = "rdfrdream"
    props.title = "基于 ESP32-C3 与 MLX90640 的热成像相机系统项目报告"
    props.subject = "成果提交报告"
    return doc


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = False
    return p


def add_para(doc: Document, text: str, size: float = 10.5, bold=False, color=BODY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_note(doc: Document, title: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    cell = table.cell(0, 0)
    shade_cell(cell, NOTE_FILL)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=ACCENT_DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, image_path: Path, caption: str, width: float):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(8)
    r = cp.add_run(caption)
    set_run_font(r, size=9.3, color=MUTED)


def add_small_table(doc: Document, rows: list[tuple[str, str, str]]):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    headers = ("模块", "实现内容", "对应文件")
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        shade_cell(cell, LIGHT_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        r = p.add_run(header)
        set_run_font(r, size=10, bold=True, color=ACCENT_DARK)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            r = p.add_run(value)
            set_run_font(r, size=9.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("基于 ESP32-C3 与 MLX90640 的\n热成像相机系统项目报告")
    set_run_font(r, size=24, bold=True, color=ACCENT_DARK)
    paragraph_border_bottom(p)

    sub = doc.add_paragraph()
    r = sub.add_run("成果提交版 | PDF 5000 字以内 | rdfrdream")
    set_run_font(r, size=12.5, color=MUTED)

    add_note(
        doc,
        "摘要",
        "本项目设计并实现了一套基于 ESP32-C3 与 MLX90640 的热成像相机系统。系统使用 I2C 读取 MLX90640 32 x 24 红外阵列温度数据，利用 ESP32-C3 的 WiFi 能力建立局域网 Web 服务，并通过浏览器网页完成热图显示、色温条、数据刷新与 CSV 导出。由于样机阶段采用裸传感器、面包板和杜邦线连接，真实采集链路仍存在稳定性验证问题，因此展示环节结合前期实验视频和网页端演示数据说明系统流程。",
    )


def build_report():
    assets = prepare_assets()
    doc = setup_doc()
    add_title_page(doc)

    add_heading(doc, "1. 背景与意义", 1)
    add_para(
        doc,
        "热成像可以把不可见的红外辐射转换为可观察的温度分布，在电路板发热检测、设备状态观察、教学演示和低成本测温实验中具有直观价值。传统热像仪价格较高，而 ESP32-C3 自带 WiFi/BLE，MLX90640 具有 32 x 24 阵列输出能力，二者组合适合构建低成本、可联网、可展示的热成像实验系统。",
    )
    add_para(
        doc,
        "本项目的课程意义在于把微机接口技术中的 I2C 通信、嵌入式软件、串口调试、无线网络和 Web 可视化结合起来，从而形成一个完整的数据采集与显示链路。",
    )
    add_figure(doc, assets["flow"], "图 1 系统总体数据链路", width=6.6)

    add_heading(doc, "2. 方法描述", 1)
    add_para(
        doc,
        "系统采用分层设计。硬件层由 ESP32-C3、MLX90640、上拉电阻、去耦电容和面包板组成；通信层使用 I2C 总线，其中 SDA 接 GPIO4，SCL 接 GPIO5；主控层完成 MLX90640 初始化、帧数据读取、温度计算和缓存；网络层由 ESP32-C3 建立 WiFi AP 与 HTTP Server；显示层由浏览器网页完成热图绘制和数据导出。",
    )
    add_small_table(
        doc,
        [
            ("采集模块", "I2C 初始化、MLX90640 参数读取、温度帧计算", "src/main.c, lib/MLX90640"),
            ("网络模块", "WiFi AP、HTTP Server、JSON 数据接口", "src/main.c"),
            ("显示模块", "Canvas 热图、色温条、导入数据、CSV 下载", "data/web_script.js"),
            ("演示模块", "生成平滑移动热源，用于无硬件演示", "tools/mock_thermal_server.py"),
        ],
    )
    add_figure(doc, assets["wiring"], "图 2 面包板样机与接线过程", width=4.4)

    add_heading(doc, "3. 关键代码与截图", 1)
    add_para(
        doc,
        "关键代码主要集中在三部分：第一，MLX90640 初始化与参数提取；第二，循环读取帧数据并计算温度矩阵；第三，将温度数组封装为 JSON 供网页端请求。网页端使用 JavaScript 定时获取数据，并将温度值映射为热图颜色。",
    )
    add_figure(doc, assets["code_init"], "图 3 MLX90640 初始化与参数读取代码", width=5.8)
    add_figure(doc, assets["code_api"], "图 4 热图数据接口代码：返回 JSON 温度数组", width=5.8)

    add_heading(doc, "4. 实验结果", 1)
    add_para(
        doc,
        "软件链路已经完成并可运行：PlatformIO 工程能够编译，ESP32-C3 固件和 SPIFFS 网页资源能够烧录，浏览器可进入热成像网页，网页端能够显示热图、色温条，并支持导入演示数据与导出 CSV。网页调试数据采用缓慢移动的热源模型，使画面变化更接近真实采样过程，便于录制演示视频。",
    )
    add_figure(doc, assets["web_debug"], "图 5 网页热图、色温条与 CSV 功能演示效果", width=6.4)
    add_note(
        doc,
        "实地验证说明",
        "老师要求现场实地验证时，项目已完成实物搭建和网页链路验证，但 MLX90640 裸传感器与面包板连接存在松动和 I2C 稳定性问题，导致真实采集数据不够连续。因此展示时采用前期实验视频和网页端演示数据补充说明核心功能。该问题主要集中在硬件连接稳定性，而不是网页显示或软件接口逻辑。",
    )

    add_heading(doc, "5. 讨论与改进方向", 1)
    add_para(
        doc,
        "本项目证明了 ESP32-C3 可以同时承担主控、无线通信和小型 Web Server 三个角色，比 C51 更适合联网显示，比普通 STM32 方案更少依赖外接 WiFi 模块。现阶段主要不足在于裸传感器接线和供电抗干扰能力不足，导致实时采集稳定性仍需验证。",
    )
    add_bullets(
        doc,
        [
            "硬件改进：使用 MLX90640 转接板、焊接固定或 PCB，增加外壳与传感器固定结构。",
            "电气改进：优化 SDA/SCL 上拉电阻、供电去耦和线长，降低 I2C 通信错误。",
            "算法改进：加入异常点滤波、帧间平滑、温度校准和热点追踪。",
            "功能拓展：温度极值触发 LED 或报警器，加入阈值报警、远程访问和云端记录。",
            "应用拓展：结合红外识别算法，发展为电路板发热检测或自动异常识别装置。",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "6. 提交材料说明", 1)
    add_para(
        doc,
        "本项目已整理完整源代码、README 运行说明、项目报告、答辩 PPT 和演示材料。GitHub 仓库中 README.md 给出了编译、烧录、SPIFFS 上传、真实硬件访问、本地演示和调试接口说明。项目报告用于成果提交，答辩 PPT 用于课堂展示。",
    )
    add_bullets(
        doc,
        [
            "完整源代码：src、data、lib、tools、platformio.ini、partitions.csv。",
            "项目报告：deliverables/基于ESP32-C3与MLX90640的热成像相机系统项目报告_提交版.pdf。",
            "答辩 PPT：deliverables/基于ESP32-C3与MLX90640的热成像相机系统设计_答辩PPT.pptx。",
            "仓库地址：https://github.com/1909899270h-spec/esp32c3-mlx90640-thermal-camera。",
        ],
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_report()
