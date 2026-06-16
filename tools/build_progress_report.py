from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"H:\ircam_pio_c3")
OUT_DIR = ROOT / "deliverables"
ASSET_DIR = ROOT / "outputs" / "manual-20260605-ircam" / "presentations" / "ircam-progress" / "assets"
OUT = OUT_DIR / "热成像相机项目阶段性报告.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_run(run, size=10.5, bold=False, color="1F2937") -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc: Document, text: str = "", style: str | None = None, bold: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    style_run(run, bold=bold)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    style_run(run, size=10.5)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    style_run(run, size=10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        set_cell_shading(hdr[i], "DDEBFF")
        set_cell_text(hdr[i], text, bold=True, color="0B2545")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_image(doc: Document, image_name: str, caption: str, width_cm: float = 14.5) -> None:
    path = ASSET_DIR / image_name
    if not path.exists():
        add_para(doc, f"[图片缺失：{image_name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    style_run(run, size=9, color="6B7280")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, "0B2545"),
        ("Heading 2", 13, "1F4D78"),
        ("Heading 3", 11.5, "374151"),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("ESP32-C3 + MLX90640 热成像相机阶段性报告")
    style_run(run, size=8.5, color="6B7280")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("基于 ESP32-C3 与 MLX90640 的热成像相机项目阶段性报告")
    style_run(run, size=20, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("微机接口与技术课程大作业 | 当前版本：软件与网页端验证完成，硬件传感器数据链路待完善")
    style_run(run, size=10.5, color="4B5563")

    doc.add_paragraph()
    add_table(
        doc,
        ["项目要素", "当前内容"],
        [
            ["主控平台", "ESP32-C3 开发板（CORE-ESP32-C3 / CH343 串口，COM5）"],
            ["传感器", "MLX90640ESF-BAA/BAB 裸传感器，I2C 地址 0x33"],
            ["开发环境", "VS Code + PlatformIO + ESP-IDF 5.4.0"],
            ["当前状态", "固件与 SPIFFS 已烧录，热点网页已打开，热成像数据仍待硬件完善"],
        ],
        [4.0, 11.2],
    )

    doc.add_heading("1. 项目目标与总体思路", level=1)
    add_para(
        doc,
        "本项目计划实现一个可通过浏览器访问的低成本热成像相机。ESP32-C3 负责采集 MLX90640 红外阵列传感器数据，"
        "并在本地开启 Wi-Fi 热点与 HTTP 服务；手机或电脑连接热点后，通过网页查看热成像画面和温度范围。"
    )
    add_bullet(doc, "硬件侧：ESP32-C3 与 MLX90640 通过 I2C 总线连接，传感器使用 3.3V 供电。")
    add_bullet(doc, "软件侧：使用 ESP-IDF 工程，完成 I2C、SPIFFS、Wi-Fi AP、HTTP 服务和网页端显示逻辑。")
    add_bullet(doc, "展示侧：网页端已能打开控制界面，后续传感器通讯稳定后即可显示热图数据。")

    doc.add_heading("2. 当前已完成的软件工作", level=1)
    add_table(
        doc,
        ["工作项", "完成情况", "说明"],
        [
            ["工程迁移", "已完成", r"整理到 H:\ircam_pio_c3，避开中文/空格路径导致的工具链问题。"],
            ["PlatformIO 环境", "已完成", "安装并缓存 espressif32@6.10.0，框架为 ESP-IDF 5.4.0。"],
            ["固件编译", "已完成", "firmware.bin 编译成功，Flash 使用约 88.1%。"],
            ["SPIFFS 构建", "已完成", "web_script.js 已打包进入 spiffs.bin，大小 1MB。"],
            ["烧录验证", "已完成", "COM5 识别为 CH343，bootloader、分区表、主固件与 SPIFFS 均写入并校验通过。"],
            ["网页访问", "已完成", "浏览器已能访问控制页面，说明热点、HTTP 服务和网页资源已可用。"],
        ],
        [3.0, 2.4, 9.8],
    )

    doc.add_heading("3. 关键配置与烧录记录", level=1)
    add_table(
        doc,
        ["配置项", "当前值"],
        [
            ["I2C SDA", "GPIO4 / IO04"],
            ["I2C SCL", "GPIO5 / IO05"],
            ["I2C 频率", "400 kHz"],
            ["MLX90640 地址", "0x33"],
            ["AP 热点命名", "IRCAM-XXXX，根据 MAC 后两字节生成"],
            ["网页入口", "http://192.168.4.1"],
            ["数据接口", "http://192.168.4.1/thermal-data"],
        ],
        [4.0, 11.2],
    )
    add_para(
        doc,
        "烧录阶段已读取到 ESP32-C3 芯片信息，MAC 为 88:56:a6:e8:b9:c8。主固件写入地址为 0x10000，"
        "SPIFFS 文件系统写入地址为 0x110000，均通过 hash 校验。"
    )

    add_image(doc, "webpage.jpg", "图 1：当前网页端已能打开，但传感器数据尚未进入，温度显示仍为 --°C", 10.0)

    doc.add_heading("4. 当前硬件连接方案", level=1)
    add_para(doc, "根据开发板背面丝印与已同步的软件配置，当前推荐接线如下。注意 MLX90640 为裸传感器，必须避免接入 5V。")
    add_table(
        doc,
        ["ESP32-C3 引脚", "MLX90640 引脚", "作用"],
        [
            ["3.3V", "VDD", "传感器 3.3V 供电"],
            ["GND", "GND", "公共地"],
            ["IO04", "SDA", "I2C 数据线"],
            ["IO05", "SCL", "I2C 时钟线"],
        ],
        [4.0, 4.0, 7.2],
    )
    add_para(doc, "MLX90640 底面朝上且定位小凸起在上方时，引脚方向为：左上 SCL，右上 SDA，左下 GND，右下 VDD。")
    add_image(doc, "esp32_back.jpg", "图 2：ESP32-C3 背面丝印，用于确认 IO04、IO05、3.3V 与 GND", 14.0)
    add_image(doc, "wiring_current.jpg", "图 3：当前临时接线状态，后续需补上 I2C 上拉电阻与供电电容", 14.0)

    doc.add_heading("5. 当前硬件待完善项", level=1)
    add_table(
        doc,
        ["待完善项", "推荐做法", "作用"],
        [
            ["I2C 上拉电阻", "SDA 到 3.3V 接 5kΩ，SCL 到 3.3V 接 5kΩ", "保证 I2C 总线高电平稳定，减少通信失败。"],
            ["供电电容", "VDD 与 GND 之间接 10uF；若后续有 0.1uF 可并联补充", "降低瞬态电压波动，提高传感器初始化稳定性。"],
            ["裸传感器固定", "用胶带或更稳的插接方式分隔四根母头线", "避免相邻引脚短路或接触不良。"],
            ["I2C 验证", "访问 /thermal-data 或查看串口日志", "判断 MLX90640 是否成功初始化并输出数据。"],
        ],
        [3.6, 6.2, 5.4],
    )

    doc.add_heading("6. 阶段性结论", level=1)
    add_para(
        doc,
        "目前项目已经完成软件平台搭建、ESP32-C3 固件烧录、SPIFFS 网页资源烧录和网页端访问验证。"
        "网页端黑屏和温度显示 --°C 并不代表网页失败，而是说明当前还没有拿到有效的 MLX90640 温度数据。"
    )
    add_para(
        doc,
        "下一阶段工作的重点是补齐裸传感器所需的 I2C 上拉和供电滤波，确认 IO04/IO05 与 SDA/SCL 的连接稳定，"
        "随后通过串口日志和 /thermal-data 接口验证传感器数据链路。"
    )

    doc.add_heading("7. 后续计划", level=1)
    for item in [
        "断电后补接两个 5kΩ 上拉电阻和一个 10uF 供电电容。",
        "重新上电，观察热点与网页是否正常出现。",
        "访问 /thermal-data，判断是否返回 768 个温度点或 JSON 数据。",
        "若仍失败，查看串口日志中的 MLX90640 EEPROM/I2C 错误信息，逐项排查 SDA、SCL、VDD、GND。",
        "硬件数据链路打通后，补充最终热成像截图、性能说明和课堂展示视频/照片。",
    ]:
        add_number(doc, item)

    doc.add_page_break()
    doc.add_heading("附录：当前产物路径", level=1)
    add_table(
        doc,
        ["产物", "路径"],
        [
            ["工程目录", r"H:\ircam_pio_c3"],
            ["VS Code 入口", r"H:\ircam_pio_c3\Open_PlatformIO_VSCode.cmd"],
            ["主固件", r"H:\ircam_pio_c3\.pio\build\esp32-c3-devkitm-1\firmware.bin"],
            ["SPIFFS 镜像", r"H:\ircam_pio_c3\.pio\build\esp32-c3-devkitm-1\spiffs.bin"],
            ["报告输出", str(OUT)],
        ],
        [4.0, 11.2],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
